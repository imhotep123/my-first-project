#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
規約比較ビューワー v2
- 各パネル独立した条文番号検索
- 条文名（第X条直前の（...）行）も含めて表示
- テキストファイル出力機能
- ESC キーで両パネル一括検索
"""

from __future__ import annotations

import re
import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox


# ─────────────────────────────────────────
#  依存ライブラリ確認
# ─────────────────────────────────────────

def _ensure_docx() -> None:
    try:
        import docx  # noqa: F401
    except ImportError:
        root_tmp = tk.Tk()
        root_tmp.withdraw()
        yes = messagebox.askyesno(
            "python-docx が見つかりません",
            "Word ファイルを読み込むには python-docx が必要です。\n"
            "今すぐインストールしますか？\n\n  pip install python-docx",
        )
        root_tmp.destroy()
        if yes:
            import subprocess
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            messagebox.showinfo("完了", "インストールしました。アプリを再起動してください。")
        sys.exit(0)


_ensure_docx()

from docx import Document  # noqa: E402


# ─────────────────────────────────────────
#  数字正規化
# ─────────────────────────────────────────

_FULLWIDTH = str.maketrans("０１２３４５６７８９", "0123456789")
_KANJI_UNIT = {
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9,
}


def normalize_num(s: str) -> str:
    """全角数字 / 漢数字 → 半角アラビア数字へ変換"""
    if not s:
        return s
    s = s.translate(_FULLWIDTH)
    if re.fullmatch(r"[一二三四五六七八九十百千]+", s):
        n = cur = 0
        for ch in s:
            if ch == "十":
                n += (cur or 1) * 10; cur = 0
            elif ch == "百":
                n += (cur or 1) * 100; cur = 0
            elif ch == "千":
                n += (cur or 1) * 1000; cur = 0
            else:
                cur = _KANJI_UNIT.get(ch, 0)
        return str(n + cur)
    return s


# ─────────────────────────────────────────
#  Word 解析
# ─────────────────────────────────────────

# 条文番号行パターン（「第X条」「第X条のY」）
_ART_RE = re.compile(
    r"^[\s　]*第\s*(?P<art>[０-９\d一二三四五六七八九十百千]+)\s*条"
    r"(?:\s*の\s*(?P<sub>[０-９\d一二三四五六七八九十百千]+))?"
)

# 条文名行パターン（「（目的）」「（定義）」など、行全体が括弧で囲まれているもの）
_NAME_RE = re.compile(r"^[\s　]*[（(].+[）)][\s　]*$")


def _make_key(art_raw: str, sub_raw: str | None) -> str:
    art = normalize_num(art_raw)
    return f"第{art}条の{normalize_num(sub_raw)}" if sub_raw else f"第{art}条"


def parse_docx(path: str) -> dict[str, str]:
    """Word ファイルを解析し {条文キー: 本文テキスト} を返す。
    条文番号の直前にある「（条文名）」行も先頭に含める。
    """
    doc = Document(path)
    lines: list[str] = []

    for para in doc.paragraphs:
        lines.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        lines.append(para.text)

    articles: dict[str, str] = {}
    cur_key: str | None = None
    buf: list[str] = []
    prev_nonempty: str = ""  # 直前の非空行（条文名の候補）

    def flush() -> None:
        if cur_key is not None:
            articles[cur_key] = "\n".join(buf).strip()

    for raw in lines:
        stripped = raw.strip()
        m = _ART_RE.match(raw)
        if m:
            # 直前の非空行が（条文名）なら、それは「次の条文の名前」
            # → 現在の条文 buf の末尾から除去してから flush し、
            #   新しい条文の先頭に付加する
            article_name: str | None = None
            if prev_nonempty and _NAME_RE.match(prev_nonempty):
                for i in range(len(buf) - 1, -1, -1):
                    if buf[i] == prev_nonempty:
                        article_name = buf.pop(i)
                        # 末尾に残った空行も除去
                        while buf and not buf[-1]:
                            buf.pop()
                        break

            flush()
            cur_key = _make_key(m.group("art"), m.group("sub"))
            buf = [article_name, stripped] if article_name else [stripped]
        elif cur_key is not None:
            if stripped:
                buf.append(stripped)
            elif buf and buf[-1]:
                buf.append("")  # 段落間の空行を保持

        if stripped:
            prev_nonempty = stripped

    flush()
    return articles


def find_article(articles: dict, key: str) -> str | None:
    if not articles:
        return None
    if key in articles:
        return articles[key]
    kn = normalize_num(key)
    for k, v in articles.items():
        if normalize_num(k) == kn:
            return v
    return None


# ─────────────────────────────────────────
#  定数 / スタイル
# ─────────────────────────────────────────

F_BODY  = ("MS Gothic", 11)
F_BOLD  = ("MS Gothic", 11, "bold")
F_SMALL = ("MS Gothic",  9)

C_LEFT_H  = "#2c5f9e"
C_RIGHT_H = "#2a7a4a"
BG_LEFT   = "#fffef2"
BG_RIGHT  = "#f2fff5"


# ─────────────────────────────────────────
#  DocPanel: 1規約分の表示パネル
# ─────────────────────────────────────────

class DocPanel(tk.Frame):
    """条文検索バー＋テキスト表示エリアをひとまとめにしたパネル"""

    def __init__(
        self,
        parent: tk.Widget,
        base_title: str,
        bg_color: str,
        header_color: str,
        **kwargs,
    ) -> None:
        super().__init__(parent, **kwargs)
        self._base_title  = base_title
        self._bg_color    = bg_color
        self._header_color = header_color

        self.articles: dict[str, str] = {}
        self.doc_name: str = ""          # 拡張子なしのファイル名
        self._current_key: str | None = None
        self._has_content: bool = False  # 条文が見つかったか

        self._build()

    # ── UI 構築 ──────────────────────────

    def _build(self) -> None:
        # ヘッダー（ファイル名表示）
        self._header = tk.Label(
            self, text=self._base_title, font=F_BOLD,
            bg=self._header_color, fg="white", padx=12, pady=5, anchor=tk.W,
        )
        self._header.pack(fill=tk.X)

        # 検索バー
        sf = tk.Frame(self, bg="#ebebeb", padx=8, pady=5)
        sf.pack(fill=tk.X)

        tk.Label(sf, text="第", font=("MS Gothic", 11), bg="#ebebeb").pack(side=tk.LEFT)
        self._art_var = tk.StringVar()
        self._art_ent = tk.Entry(sf, textvariable=self._art_var,
                                  font=("MS Gothic", 11), width=6)
        self._art_ent.pack(side=tk.LEFT, padx=(2, 0))

        tk.Label(sf, text="条の", font=("MS Gothic", 11), bg="#ebebeb").pack(side=tk.LEFT)
        self._sub_var = tk.StringVar()
        self._sub_ent = tk.Entry(sf, textvariable=self._sub_var,
                                  font=("MS Gothic", 11), width=5)
        self._sub_ent.pack(side=tk.LEFT, padx=(2, 6))

        tk.Label(sf, text="（枝番なければそのままEnter）",
                 font=("MS Gothic", 8), fg="#888", bg="#ebebeb").pack(side=tk.LEFT)

        tk.Button(
            sf, text="検索", command=self._on_search,
            padx=10, pady=2, bg="#3572b8", fg="white", relief=tk.GROOVE,
        ).pack(side=tk.RIGHT, padx=4)

        self._key_lbl = tk.Label(sf, text="", font=("MS Gothic", 10, "bold"),
                                  bg="#ebebeb", fg="#333")
        self._key_lbl.pack(side=tk.RIGHT, padx=8)

        # Enterキーの動作
        self._art_ent.bind("<Return>", lambda _: self._sub_ent.focus_set())
        self._sub_ent.bind("<Return>", lambda _: self._on_search())

        # テキスト表示エリア
        txt_frame = tk.Frame(self)
        txt_frame.pack(fill=tk.BOTH, expand=True)

        self._txt = tk.Text(
            txt_frame, wrap=tk.WORD, font=F_BODY,
            padx=14, pady=14, bg=self._bg_color,
            state=tk.DISABLED, relief=tk.FLAT, cursor="arrow",
            selectbackground="#b0d0f0",
        )
        sy = tk.Scrollbar(txt_frame, orient=tk.VERTICAL,   command=self._txt.yview)
        sx = tk.Scrollbar(txt_frame, orient=tk.HORIZONTAL, command=self._txt.xview)
        self._txt.config(yscrollcommand=sy.set, xscrollcommand=sx.set)
        sy.pack(side=tk.RIGHT,  fill=tk.Y)
        sx.pack(side=tk.BOTTOM, fill=tk.X)
        self._txt.pack(fill=tk.BOTH, expand=True)

        # 初期プレースホルダ
        self._placeholder(f"「{self._base_title}を開く」でファイルを選択してください")

    # ── ファイルセット ────────────────────

    def load_file(self, path: str) -> None:
        """ファイルを解析してパネルをリセット"""
        articles = parse_docx(path)
        self.articles = articles
        self.doc_name = os.path.splitext(os.path.basename(path))[0]
        self._current_key = None
        self._has_content = False
        self._header.config(
            text=f"{self._base_title}  ―  {os.path.basename(path)}"
        )
        self._key_lbl.config(text="")
        self._placeholder(
            f"読み込み完了 ({len(articles)} 条文)\n\n条文番号を入力して検索してください"
        )

    # ── 検索 ──────────────────────────────

    def _on_search(self) -> None:
        """検索バーから検索実行"""
        art = self._art_var.get().strip()
        sub = self._sub_var.get().strip()
        if not art:
            messagebox.showwarning("入力エラー", "条文番号を入力してください", parent=self)
            self._art_ent.focus_set()
            return
        key = (
            f"第{normalize_num(art)}条の{normalize_num(sub)}"
            if sub else f"第{normalize_num(art)}条"
        )
        self.show_key(key)

    def show_key(self, key: str) -> None:
        """外部からキーを指定して表示（ESC一括検索で使用）"""
        # 検索バーの表示も更新
        m = re.match(r"第(\d+)条(?:の(\d+))?", key)
        if m:
            self._art_var.set(m.group(1))
            self._sub_var.set(m.group(2) or "")
        self._render_key(key)

    def _render_key(self, key: str) -> None:
        content = find_article(self.articles, key)
        self._current_key = key
        self._has_content = content is not None
        self._key_lbl.config(text=f"表示: {key}" if content else f"未検出: {key}")
        self._render(content, key)

    def _render(self, content: str | None, key: str) -> None:
        self._txt.config(state=tk.NORMAL)
        self._txt.delete("1.0", tk.END)

        self._txt.tag_config("name", font=F_BOLD, foreground="#555577")
        self._txt.tag_config("h",    font=F_BOLD, foreground="#111111")
        self._txt.tag_config("body", font=F_BODY, foreground="#222222")
        self._txt.tag_config("err",  font=F_BODY, foreground="#cc2222")
        self._txt.tag_config("hint", font=F_SMALL, foreground="#999999")
        self._txt.tag_config("ph",   font=F_BODY, foreground="#aaaaaa")

        if content:
            lines = content.splitlines()
            idx = 0
            # 先頭行が条文名「（...）」かどうか判定
            if lines and _NAME_RE.match(lines[0]) and len(lines) > 1:
                self._txt.insert(tk.END, lines[0] + "\n", "name")
                self._txt.insert(tk.END, lines[1] + "\n\n", "h")
                idx = 2
            else:
                self._txt.insert(tk.END, lines[0] + "\n\n", "h")
                idx = 1
            if idx < len(lines):
                self._txt.insert(tk.END, "\n".join(lines[idx:]), "body")

        elif self.articles:
            self._txt.insert(tk.END, f"「{key}」は見つかりませんでした。\n\n", "err")

            def _sort_key(k: str) -> int:
                m2 = re.search(r"\d+", k)
                return int(m2.group()) if m2 else 0

            sample = sorted(self.articles.keys(), key=_sort_key)[:15]
            self._txt.insert(tk.END, "【存在する条文（先頭 15 件）】\n", "hint")
            for k in sample:
                self._txt.insert(tk.END, f"    {k}\n", "hint")
        else:
            self._txt.insert(tk.END, "ファイルが読み込まれていません", "ph")

        self._txt.config(state=tk.DISABLED)

    def _placeholder(self, msg: str) -> None:
        self._txt.config(state=tk.NORMAL)
        self._txt.delete("1.0", tk.END)
        self._txt.tag_config("ph", foreground="#aaa", font=F_BODY)
        self._txt.insert(tk.END, f"\n\n  {msg}", "ph")
        self._txt.config(state=tk.DISABLED)

    # ── テキスト出力用 ────────────────────

    def get_export_info(self) -> tuple[str | None, str | None, str | None]:
        """(doc_name, current_key, content_text) を返す。未検索・未検出は None"""
        if not self._has_content or self._current_key is None:
            return None, None, None
        content = find_article(self.articles, self._current_key)
        return self.doc_name, self._current_key, content


# ─────────────────────────────────────────
#  メインウィンドウ
# ─────────────────────────────────────────

class CompareViewer:

    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self.root.title("規約比較ビューワー")
        self.root.geometry("1440x880")
        self._build_ui()
        self.root.bind("<Escape>", lambda _: self._unified_search())

    def _build_ui(self) -> None:
        # ── ツールバー ──
        bar = tk.Frame(self.root, bg="#e0e0e0", padx=8, pady=6)
        bar.pack(fill=tk.X)

        def _btn(text, cmd, **kw):
            return tk.Button(bar, text=text, command=cmd,
                             padx=9, pady=3, relief=tk.GROOVE, **kw)

        _btn("比較元規約を開く",   self._load1).pack(side=tk.LEFT, padx=3)
        _btn("比較対象規約を開く", self._load2).pack(side=tk.LEFT, padx=3)

        tk.Frame(bar, width=1, bg="#bbb").pack(side=tk.LEFT, fill=tk.Y, padx=10, pady=2)

        _btn("両パネル一括検索  [ESC]", self._unified_search,
             bg="#3572b8", fg="white").pack(side=tk.LEFT, padx=3)

        tk.Frame(bar, width=1, bg="#bbb").pack(side=tk.LEFT, fill=tk.Y, padx=10, pady=2)

        _btn("テキスト出力", self._export_text,
             bg="#4a8a30", fg="white").pack(side=tk.LEFT, padx=3)

        self._status = tk.StringVar(value="ファイルを読み込んでください")
        tk.Label(bar, textvariable=self._status, bg="#e0e0e0",
                 fg="#555", font=F_SMALL).pack(side=tk.RIGHT, padx=10)

        # ── 分割ペイン ──
        pw = tk.PanedWindow(self.root, orient=tk.HORIZONTAL,
                             sashwidth=5, sashrelief=tk.FLAT, bg="#bbb")
        pw.pack(fill=tk.BOTH, expand=True, padx=4, pady=(0, 4))

        self._panel1 = DocPanel(pw, "比較元規約",   BG_LEFT,  C_LEFT_H)
        pw.add(self._panel1, minsize=400, stretch="always")

        self._panel2 = DocPanel(pw, "比較対象規約", BG_RIGHT, C_RIGHT_H)
        pw.add(self._panel2, minsize=400, stretch="always")

        # ── ステータスバー ──
        sb = tk.Frame(self.root, bg="#d0d0d0", height=22)
        sb.pack(fill=tk.X, side=tk.BOTTOM)
        sb.pack_propagate(False)
        tk.Label(sb, textvariable=self._status, bg="#d0d0d0",
                 font=F_SMALL, anchor=tk.W).pack(side=tk.LEFT, padx=8, pady=2)

    # ── ファイル読み込み ────────────────

    def _load1(self) -> None:
        path = filedialog.askopenfilename(
            title="比較元規約ファイルを選択",
            filetypes=[("Word 文書", "*.docx"), ("すべてのファイル", "*.*")],
        )
        if not path:
            return
        try:
            self._panel1.load_file(path)
            self._status.set(
                f"比較元規約 読み込み完了  ({len(self._panel1.articles)} 条文検出)"
            )
        except Exception as exc:
            messagebox.showerror("読み込みエラー", str(exc))

    def _load2(self) -> None:
        path = filedialog.askopenfilename(
            title="比較対象規約ファイルを選択",
            filetypes=[("Word 文書", "*.docx"), ("すべてのファイル", "*.*")],
        )
        if not path:
            return
        try:
            self._panel2.load_file(path)
            self._status.set(
                f"比較対象規約 読み込み完了  ({len(self._panel2.articles)} 条文検出)"
            )
        except Exception as exc:
            messagebox.showerror("読み込みエラー", str(exc))

    # ── 一括検索（ESC） ─────────────────

    def _unified_search(self) -> None:
        """ESC キー: 両パネルに同じ条文番号で検索"""
        if not self._panel1.articles and not self._panel2.articles:
            messagebox.showinfo("お知らせ", "先にファイルを読み込んでください")
            return
        dlg = SearchDialog(self.root)
        self.root.wait_window(dlg)
        if dlg.result:
            art, sub = dlg.result
            key = (
                f"第{normalize_num(art)}条の{normalize_num(sub)}"
                if sub else f"第{normalize_num(art)}条"
            )
            self._panel1.show_key(key)
            self._panel2.show_key(key)
            self._status.set(f"両パネルで {key} を検索しました")

    # ── テキスト出力 ───────────────────

    def _export_text(self) -> None:
        """表示中の条文をテキストファイルに出力（条文番号のフォルダを作成して保存）"""
        name1, key1, text1 = self._panel1.get_export_info()
        name2, key2, text2 = self._panel2.get_export_info()

        if not text1 and not text2:
            messagebox.showinfo(
                "出力できません",
                "出力する条文がありません。\n先に条文を検索してください。",
            )
            return

        save_dir = filedialog.askdirectory(title="保存先フォルダを選択")
        if not save_dir:
            return

        # フォルダ名は「比較元規約」の条文番号を使用（なければ比較対象規約の条文番号）
        folder_key = key1 or key2
        folder_key_safe = re.sub(r'[\\/:*?"<>|]', "_", folder_key)
        folder_path = os.path.join(save_dir, folder_key_safe)

        try:
            os.makedirs(folder_path, exist_ok=True)
        except OSError as e:
            messagebox.showerror("フォルダ作成エラー", f"フォルダを作成できませんでした:\n{e}")
            return

        saved: list[str] = []
        errors: list[str] = []

        for doc_name, key, text in [
            (name1, key1, text1),
            (name2, key2, text2),
        ]:
            if not text or not doc_name or not key:
                continue
            # ファイル名に使えない文字を除去
            key_safe = re.sub(r'[\\/:*?"<>|]', "_", key)
            fname = f"{doc_name}_{key_safe}.txt"
            fpath = os.path.join(folder_path, fname)
            try:
                with open(fpath, "w", encoding="utf-8") as f:
                    f.write(text)
                saved.append(fname)
            except OSError as e:
                errors.append(f"{fname}: {e}")

        if saved:
            msg = f"フォルダ「{folder_key_safe}」に以下のファイルを保存しました:\n\n" + "\n".join(saved)
            if errors:
                msg += "\n\nエラー:\n" + "\n".join(errors)
            messagebox.showinfo("出力完了", msg)
            self._status.set(f"テキスト出力完了: {folder_key_safe}/ → {', '.join(saved)}")
        else:
            messagebox.showerror("出力失敗", "\n".join(errors))


# ─────────────────────────────────────────
#  一括検索ダイアログ（ESC キー用）
# ─────────────────────────────────────────

class SearchDialog(tk.Toplevel):

    def __init__(self, parent: tk.Tk) -> None:
        super().__init__(parent)
        self.result: tuple[str, str] | None = None
        self.title("両パネル一括検索")
        self.geometry("420x260")
        self.resizable(False, False)
        self.transient(parent)
        self.grab_set()

        self.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width()  - 420) // 2
        y = parent.winfo_y() + (parent.winfo_height() - 260) // 2
        self.geometry(f"+{x}+{y}")

        self._build()
        self.bind("<Escape>", self._cancel)
        self._art_ent.focus_set()

    def _build(self) -> None:
        outer = tk.Frame(self, padx=30, pady=24)
        outer.pack(fill=tk.BOTH, expand=True)

        # 条文番号
        tk.Label(outer, text="表示する条文番号を入力してください。",
                 font=("MS Gothic", 11)).grid(
            row=0, column=0, columnspan=4, sticky=tk.W, pady=(0, 6))

        tk.Label(outer, text="第", font=("MS Gothic", 13)).grid(
            row=1, column=0, sticky=tk.W)
        self._art_var = tk.StringVar()
        self._art_ent = tk.Entry(outer, textvariable=self._art_var,
                                  font=("MS Gothic", 13), width=8)
        self._art_ent.grid(row=1, column=1, sticky=tk.W, padx=4)
        tk.Label(outer, text="条", font=("MS Gothic", 13)).grid(
            row=1, column=2, sticky=tk.W)

        # 枝番号
        tk.Label(outer, text="枝番号を入力してください。（なければそのままEnterを押す）",
                 font=("MS Gothic", 11)).grid(
            row=2, column=0, columnspan=4, sticky=tk.W, pady=(20, 6))

        tk.Label(outer, text="の", font=("MS Gothic", 13)).grid(
            row=3, column=0, sticky=tk.W)
        self._sub_var = tk.StringVar()
        self._sub_ent = tk.Entry(outer, textvariable=self._sub_var,
                                  font=("MS Gothic", 13), width=8)
        self._sub_ent.grid(row=3, column=1, sticky=tk.W, padx=4)
        tk.Label(outer, text="（省略可）", font=("MS Gothic", 9), fg="#999").grid(
            row=3, column=2, columnspan=2, sticky=tk.W)

        # ボタン
        bf = tk.Frame(outer)
        bf.grid(row=4, column=0, columnspan=4, sticky=tk.E, pady=(24, 0))
        tk.Button(bf, text="キャンセル", command=self._cancel,
                  padx=10, pady=4).pack(side=tk.RIGHT, padx=(8, 0))
        tk.Button(bf, text="検索", command=self._ok,
                  padx=16, pady=4, bg="#3572b8", fg="white",
                  relief=tk.GROOVE).pack(side=tk.RIGHT)

        self._art_ent.bind("<Return>", lambda _: self._sub_ent.focus_set())
        self._sub_ent.bind("<Return>", lambda _: self._ok())

    def _ok(self) -> None:
        art = self._art_var.get().strip()
        sub = self._sub_var.get().strip()
        if not art:
            messagebox.showwarning("入力エラー", "条文番号を入力してください", parent=self)
            self._art_ent.focus_set()
            return
        self.result = (art, sub)
        self.destroy()

    def _cancel(self, _event=None) -> None:
        self.destroy()


# ─────────────────────────────────────────
#  エントリーポイント
# ─────────────────────────────────────────

if __name__ == "__main__":
    root = tk.Tk()
    CompareViewer(root)
    root.mainloop()
