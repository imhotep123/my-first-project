# -*- coding: utf-8 -*-
"""
Comparison Chart.py

文書比較（compare_word）→ 新旧対照表作成（comparison_table）→ 整列処理（align_table）
を連続実行する完全パイプラインプログラム。

処理の流れ:
  1. compare_word     : 旧規約・改正案を比較し、消し線（赤）・アンダー（青）文書を生成
  2. comparison_table : 消し線・アンダー文書とコメントから新旧対照表を生成（中間ファイル）
  3. align_table      : 新旧対照表を条番号ごとに整列して最終ファイルを出力

入力（ダイアログで選択）:
  ① 旧規約（原文 ─ 書式なし）
  ② 改正案（原文 ─ 書式なし）
  ③ コメント
  ④ 最終出力先
"""

import os
import re
import sys
import tempfile
import tkinter as tk
from tkinter import filedialog, messagebox

# サブディレクトリのモジュールをパスに追加
_here = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.join(_here, 'compare_word'))
sys.path.insert(0, os.path.join(_here, 'comparison_table'))
sys.path.insert(0, os.path.join(_here, 'align_table'))

from compare_word import compare_documents             # noqa: E402
from comparison_table import create_comparison_table  # noqa: E402

from docx import Document                             # noqa: E402
from docx.shared import Pt, Cm                        # noqa: E402
from align_table import (                             # noqa: E402
    group_cell_paragraphs,
    group_doc_paragraphs,
    insert_missing_sections,
    copy_paras_to_cell,
    add_plain_text_to_cell,
    add_shinsetu_cell,
    get_full_text,
)


def run_alignment(taiso_path, kaiseian_path, output_path):
    """align_table の整列処理を関数として実行する。

    Args:
        taiso_path   : 新旧対照表ファイルパス（comparison_table の出力）
        kaiseian_path: アンダー文書パス（compare_word の出力、条のX 補完用）
        output_path  : 整列済み出力ファイルパス
    """
    taiso_doc = Document(taiso_path)
    tbl = taiso_doc.tables[0]
    left_cell  = tbl.rows[1].cells[0]
    mid_cell   = tbl.rows[1].cells[1]
    right_cell = tbl.rows[1].cells[2]

    left_sections  = group_cell_paragraphs(left_cell)
    mid_sections   = group_cell_paragraphs(mid_cell)
    right_sections = group_cell_paragraphs(right_cell, is_right_column=True)

    kaiseian_doc      = Document(kaiseian_path)
    kaiseian_sections = group_doc_paragraphs(kaiseian_doc)

    mid_sections = insert_missing_sections(mid_sections, kaiseian_sections)

    print(f"Left sections  (旧規約): {len(left_sections)}")
    print(f"Mid sections   (改正案): {len(mid_sections)}")
    print(f"Right sections (備考)  : {len(right_sections)}")

    # 右列マーカーマップを構築
    aligned_markers_set = set()
    for m, _ in left_sections:
        if m:
            aligned_markers_set.add(m)
    for m, _ in mid_sections:
        if m:
            aligned_markers_set.add(m)

    right_map = {}
    for marker, paras in right_sections:
        if marker:
            if marker in aligned_markers_set:
                if marker in right_map:
                    right_map[marker].extend(paras)
                else:
                    right_map[marker] = list(paras)
            else:
                base = re.sub(r'の\d+$', '', marker)
                if base in aligned_markers_set:
                    if base in right_map:
                        right_map[base].extend(paras)
                    else:
                        right_map[base] = list(paras)
                else:
                    right_map[marker] = list(paras)

    # 左列・中列を条番号で整列
    left_list = left_sections
    mid_list  = mid_sections

    left_markers_set = set(m for m, _ in left_list if m)
    mid_markers_set  = set(m for m, _ in mid_list  if m)

    aligned_rows = []
    li = 0
    mi = 0

    while li < len(left_list) or mi < len(mid_list):
        if li < len(left_list) and mi < len(mid_list):
            lm = left_list[li][0]
            mm = mid_list[mi][0]

            if lm == mm:
                aligned_rows.append((left_list[li][1], mid_list[mi][1], lm))
                li += 1
                mi += 1
            elif lm is None and mm is None:
                aligned_rows.append((left_list[li][1], mid_list[mi][1], None))
                li += 1
                mi += 1
            elif mm is not None and mm not in left_markers_set:
                # 中列にのみ存在 → 新設
                aligned_rows.append((None, mid_list[mi][1], mm))
                mi += 1
            elif lm is not None and lm not in mid_markers_set:
                # 左列にのみ存在 → 削除
                aligned_rows.append((left_list[li][1], None, lm))
                li += 1
            else:
                found_in_mid = any(mid_list[k][0] == lm for k in range(mi, len(mid_list)))
                if found_in_mid:
                    aligned_rows.append((None, mid_list[mi][1], mm))
                    mi += 1
                else:
                    aligned_rows.append((left_list[li][1], None, lm))
                    li += 1
        elif li < len(left_list):
            aligned_rows.append((left_list[li][1], None, left_list[li][0]))
            li += 1
        else:
            aligned_rows.append((None, mid_list[mi][1], mid_list[mi][0]))
            mi += 1

    new_count = sum(1 for l, _, _ in aligned_rows if l is None)
    del_count = sum(1 for _, m, _ in aligned_rows if m is None)
    print(f"\n=== Total aligned rows: {len(aligned_rows)} "
          f"(新設: {new_count}, 削除: {del_count}) ===")

    # 出力文書を作成
    new_doc = Document()

    for section in new_doc.sections:
        section.orientation    = 1
        section.page_width     = Cm(29.7)
        section.page_height    = Cm(21.0)
        section.top_margin     = Cm(1.5)
        section.bottom_margin  = Cm(1.5)
        section.left_margin    = Cm(1.5)
        section.right_margin   = Cm(1.5)

    title = new_doc.add_paragraph('新旧対照表')
    title.alignment = 1
    for run in title.runs:
        run.bold       = True
        run.font.size  = Pt(14)

    new_table = new_doc.add_table(rows=1, cols=3)
    new_table.style = 'Table Grid'

    header = new_table.rows[0]
    for idx, hdr_text in enumerate(['現行規約', '改正案', '備考']):
        cell = header.cells[idx]
        cell.text = ''
        p   = cell.paragraphs[0]
        run = p.add_run(hdr_text)
        run.bold      = True
        run.font.size = Pt(10)
        p.alignment   = 1

    for left_paras, mid_paras, marker in aligned_rows:
        row = new_table.add_row()

        if left_paras is not None:
            copy_paras_to_cell(row.cells[0], left_paras)
        else:
            mid_count = len(mid_paras) if mid_paras else 1
            add_shinsetu_cell(row.cells[0], mid_count)

        if mid_paras is not None:
            copy_paras_to_cell(row.cells[1], mid_paras)
        else:
            add_plain_text_to_cell(row.cells[1], '（削除）')

        right_paras = right_map.get(marker) if marker else None
        if right_paras is not None:
            copy_paras_to_cell(row.cells[2], right_paras)
        else:
            add_plain_text_to_cell(row.cells[2], '')

    for row in new_table.rows:
        row.cells[0].width = Cm(12)
        row.cells[1].width = Cm(12)
        row.cells[2].width = Cm(2.5)

    new_doc.save(output_path)
    print(f"\n整列済みファイルを保存しました: {output_path}")


def main():
    sys.stdout.reconfigure(encoding='utf-8')

    root = tk.Tk()
    root.withdraw()

    filetypes_word = [('Word文書', '*.docx *.doc'), ('すべてのファイル', '*.*')]

    # ① 旧規約（原文）
    old_path = filedialog.askopenfilename(
        title='① 旧規約（原文）のファイルを選択してください',
        filetypes=filetypes_word,
    )
    if not old_path:
        messagebox.showinfo('キャンセル', '旧規約のファイルが選択されませんでした。')
        return

    initial_dir = os.path.dirname(old_path)

    # ② 改正案（原文）
    new_path = filedialog.askopenfilename(
        title='② 改正案（原文）のファイルを選択してください',
        filetypes=filetypes_word,
        initialdir=initial_dir,
    )
    if not new_path:
        messagebox.showinfo('キャンセル', '改正案のファイルが選択されませんでした。')
        return

    # ③ コメント
    comment_path = filedialog.askopenfilename(
        title='③ コメントのファイルを選択してください',
        filetypes=filetypes_word,
        initialdir=initial_dir,
    )
    if not comment_path:
        messagebox.showinfo('キャンセル', 'コメントのファイルが選択されませんでした。')
        return

    # ④ 最終出力先
    output_path = filedialog.asksaveasfilename(
        title='④ 最終出力先を指定してください',
        defaultextension='.docx',
        filetypes=[('Word文書', '*.docx'), ('すべてのファイル', '*.*')],
        initialdir=initial_dir,
        initialfile='新旧対照表_整列済.docx',
    )
    if not output_path:
        messagebox.showinfo('キャンセル', '保存先が指定されませんでした。')
        return

    # 出力フォルダ（④ で指定したフォルダ）に比較結果3ファイルを保存
    output_dir          = os.path.dirname(os.path.abspath(output_path))
    compare_output_path = os.path.join(output_dir, '比較結果.docx')
    # compare_documents が自動生成するファイル（削除しない）:
    #   比較結果.docx / 比較結果_消し線.docx / 比較結果_アンダ.docx
    old_formatted_path  = None   # compare_word 出力: 消し線文書
    new_formatted_path  = None   # compare_word 出力: アンダー文書
    tmp_taiso_path      = None   # comparison_table の出力（中間・削除）

    try:
        # ─── ステップ 1/3: compare_word ─────────────────────────────────
        # 旧規約・改正案を比較し、消し線（赤）・アンダー（青）文書を出力フォルダに保存
        print('=' * 50)
        print('【ステップ 1/3】文書比較中（消し線・アンダー生成）...')
        print('=' * 50)
        old_formatted_path, new_formatted_path = compare_documents(
            old_path, new_path, compare_output_path
        )
        print('【ステップ 1/3】完了\n')

        # ─── ステップ 2/3: comparison_table ──────────────────────────────
        # 消し線・アンダー文書とコメントから新旧対照表を生成
        tmp_taiso_fd, tmp_taiso_path = tempfile.mkstemp(suffix='.docx', prefix='cc_taiso_')
        os.close(tmp_taiso_fd)

        print('=' * 50)
        print('【ステップ 2/3】新旧対照表を作成中...')
        print('=' * 50)
        create_comparison_table(
            old_formatted_path, new_formatted_path, comment_path, tmp_taiso_path
        )
        print('【ステップ 2/3】完了\n')

        # ─── ステップ 3/3: align_table ───────────────────────────────────
        # 新旧対照表を条番号ごとに整列して最終出力
        print('=' * 50)
        print('【ステップ 3/3】整列処理中...')
        print('=' * 50)
        run_alignment(tmp_taiso_path, new_formatted_path, output_path)
        print('【ステップ 3/3】完了')

        messagebox.showinfo(
            '完了',
            f'処理が完了しました。\n\n'
            f'比較結果:    {compare_output_path}\n'
            f'消し線:      {old_formatted_path}\n'
            f'アンダー:    {new_formatted_path}\n\n'
            f'整列済み出力: {output_path}',
        )

    except Exception as e:
        messagebox.showerror('エラー', f'処理中にエラーが発生しました:\n{e}')
        raise

    finally:
        # 比較結果3ファイル（比較結果・消し線・アンダー）は出力フォルダに残す
        # comparison_table の中間ファイルのみ削除
        if tmp_taiso_path and os.path.exists(tmp_taiso_path):
            try:
                os.remove(tmp_taiso_path)
            except Exception:
                pass


if __name__ == '__main__':
    main()
