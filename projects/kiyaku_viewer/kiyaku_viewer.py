#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
規約比較ビューワー v3
- Wordファイルのフォント・文字装飾（消し線・アンダーライン・色）をそのまま表示
- カーソルキー（↓/↑）で条文を順送り・前送り
- 各パネル独立した条文番号検索
- ESC キーで両パネル一括検索
- テキストファイル出力機能
"""

from __future__ import annotations

import re
import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox


# ─────────────────────────────────────────
#  依存ライブラリ確認
# ─────────────────────────────────────────

def _ensure_docx() -> None:
    try:
        import docx  # noqa: F401
    except ImportError:
        root_tmp = tk.Tk()
        root_tmp.withdraw()
        yes = messagebox.askyesno(
            "python-docx が見つかりません",
            "Word ファイルを読み込むには python-docx が必要です。\n"
            "今すぐインストールしますか？\n\n  pip install python-docx",
        )
        root_tmp.destroy()
        if yes:
            import subprocess
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            messagebox.showinfo("完了", "インストールしました。アプリを再起動してください。")
        sys.exit(0)


_ensure_docx()

from docx import Document        # noqa: E402
from docx.oxml.ns import qn as _qn  # noqa: E402


# ─────────────────────────────────────────
#  数字正規化
# ─────────────────────────────────────────

_FULLWIDTH = str.maketrans("０１２３４５６７８９", "0123456789")
_KANJI_UNIT = {
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9,
}


def normalize_num(s: str) -> str:
    """全角数字 / 漢数字 → 半角アラビア数字へ変換"""
    if not s:
        return s
    s = s.translate(_FULLWIDTH)
    if re.fullmatch(r"[一二三四五六七八九十百千]+", s):
        n = cur = 0
        for ch in s:
            if ch == "十":
                n += (cur or 1) * 10; cur = 0
            elif ch == "百":
                n += (cur or 1) * 100; cur = 0
            elif ch == "千":
                n += (cur or 1) * 1000; cur = 0
            else:
                cur = _KANJI_UNIT.get(ch, 0)
        return str(n + cur)
    return s


# ─────────────────────────────────────────
#  Word フォントマッピング（全角フォント名 → Windows フォント名）
# ─────────────────────────────────────────

_FONT_MAP: dict[str, str] = {
    "ＭＳ 明朝":        "MS Mincho",
    "MS 明朝":          "MS Mincho",
    "ＭＳ Ｐ明朝":      "MS PMincho",
    "MS P明朝":         "MS PMincho",
    "ＭＳ ゴシック":    "MS Gothic",
    "MS ゴシック":      "MS Gothic",
    "ＭＳ Ｐゴシック":  "MS PGothic",
    "MS Pゴシック":     "MS PGothic",
    "游明朝":           "Yu Mincho",
    "游ゴシック":       "Yu Gothic",
    "メイリオ":         "Meiryo",
    "BIZ UDゴシック":   "BIZ UDGothic",
    "BIZ UD明朝":       "BIZ UDMincho",
}


# ─────────────────────────────────────────
#  Word 解析（リッチテキスト版）
# ─────────────────────────────────────────

# 型メモ:
#   RichRun     = tuple[str, dict]           # (テキスト, 書式辞書)
#   RichPara    = list[RichRun]              # 1段落
#   RichArticle = list[RichPara]             # 1条文

# 条文番号行パターン（「第X条」「第X条のY」）
_ART_RE = re.compile(
    r"^[\s　]*第\s*(?P<art>[０-９\d一二三四五六七八九十百千]+)\s*条"
    r"(?:\s*の\s*(?P<sub>[０-９\d一二三四五六七八九十百千]+))?"
)

# 条文名行パターン（「（目的）」「（定義）」など）
_NAME_RE = re.compile(r"^[\s　]*[（(].+[）)][\s　]*$")


def _make_key(art_raw: str, sub_raw: str | None) -> str:
    art = normalize_num(art_raw)
    return f"第{art}条の{normalize_num(sub_raw)}" if sub_raw else f"第{art}条"


def _run_fmt(run) -> dict:
    """Word run から書式情報を辞書で返す"""
    f = run.font

    # 取り消し線：python-docx の API では取れないため XML を直接確認
    strike = False
    rpr = run._r.find(_qn('w:rPr'))
    if rpr is not None:
        strike = (
            rpr.find(_qn('w:strike'))  is not None or
            rpr.find(_qn('w:dstrike')) is not None
        )

    # 文字色（RGB のみ。テーマカラー等は無視）
    color: str | None = None
    try:
        rgb = f.color.rgb
        if rgb is not None:
            color = f"#{rgb}"
    except Exception:
        pass

    # フォントサイズ（ポイント単位）
    font_size: float | None = None
    try:
        if f.size is not None:
            font_size = f.size.pt
    except Exception:
        pass

    return {
        "bold":      bool(f.bold),
        "italic":    bool(f.italic),
        "underline": bool(f.underline),   # True/False/None/WD_UNDERLINE → bool
        "strike":    strike,
        "color":     color,
        "font_name": f.name,              # None = 継承
        "font_size": font_size,           # None = 継承
    }


def _para_to_rich(para) -> list[tuple[str, dict]]:
    """段落を (テキスト, 書式辞書) タプルのリストに変換"""
    runs = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        runs.append((text, _run_fmt(run)))
    return runs


def _para_text(rich_para: list[tuple[str, dict]]) -> str:
    """リッチ段落からプレーンテキストを抽出"""
    return "".join(text for text, _ in rich_para)


def _article_plain_text(article_data: list[list[tuple[str, dict]]]) -> str:
    """条文リッチデータをプレーンテキストに変換（エクスポート用）"""
    lines = [_para_text(para) for para in article_data]
    return "\n".join(lines).strip()


def _art_sort_key(k: str) -> tuple[int, int]:
    """条文キーをソート用タプルに変換 (条番号, 枝番号)"""
    m = re.match(r"第(\d+)条(?:の(\d+))?", k)
    if m:
        return (int(m.group(1)), int(m.group(2) or 0))
    return (9999, 0)


def parse_docx(path: str) -> dict[str, list[list[tuple[str, dict]]]]:
    """Word ファイルを解析し {条文キー: リッチ段落リスト} を返す。
    条文番号の直前にある「（条文名）」行も先頭に含める。
    """
    doc = Document(path)

    # 全段落を (プレーンテキスト, リッチ段落) として収集
    all_paras: list[tuple[str, list[tuple[str, dict]]]] = []

    for para in doc.paragraphs:
        rp = _para_to_rich(para)
        all_paras.append((_para_text(rp), rp))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    rp = _para_to_rich(para)
                    text = _para_text(rp)
                    if text.strip():
                        all_paras.append((text, rp))

    articles: dict[str, list[list[tuple[str, dict]]]] = {}
    cur_key:          str | None = None
    buf:              list[list[tuple[str, dict]]] = []
    prev_nonempty_text: str = ""

    def flush() -> None:
        if cur_key is not None:
            # 末尾の空段落を除去
            while buf and not _para_text(buf[-1]).strip():
                buf.pop()
            articles[cur_key] = list(buf)

    for text, rp in all_paras:
        stripped = text.strip()
        m = _ART_RE.match(text)
        if m:
            # 直前の非空段落が（条文名）なら現在の buf から取り出す
            article_name_para: list[tuple[str, dict]] | None = None
            if prev_nonempty_text and _NAME_RE.match(prev_nonempty_text):
                for i in range(len(buf) - 1, -1, -1):
                    if _para_text(buf[i]).strip() == prev_nonempty_text:
                        article_name_para = buf.pop(i)
                        while buf and not _para_text(buf[-1]).strip():
                            buf.pop()
                        break

            flush()
            cur_key = _make_key(m.group("art"), m.group("sub"))
            buf = []
            if article_name_para is not None:
                buf.append(article_name_para)
            buf.append(rp)   # 条文番号行

        elif cur_key is not None:
            if stripped:
                buf.append(rp)
            elif buf and _para_text(buf[-1]).strip():
                buf.append([])  # 段落間の空行

        if stripped:
            prev_nonempty_text = stripped

    flush()
    return articles


def find_article(
    articles: dict[str, list[list[tuple[str, dict]]]],
    key: str,
) -> list[list[tuple[str, dict]]] | None:
    if not articles:
        return None
    if key in articles:
        return articles[key]
    kn = normalize_num(key)
    for k, v in articles.items():
        if normalize_num(k) == kn:
            return v
    return None


# ─────────────────────────────────────────
#  定数 / スタイル
# ─────────────────────────────────────────

F_DEFAULT = ("MS Gothic", 11)
F_BOLD    = ("MS Gothic", 11, "bold")
F_SMALL   = ("MS Gothic",  9)

C_LEFT_H  = "#2c5f9e"
C_RIGHT_H = "#2a7a4a"
BG_LEFT   = "#fffef2"
BG_RIGHT  = "#f2fff5"


# ─────────────────────────────────────────
#  DocPanel: 1規約分の表示パネル
# ─────────────────────────────────────────

class DocPanel(tk.Frame):
    """条文検索バー＋テキスト表示エリアをひとまとめにしたパネル"""

    def __init__(
        self,
        parent: tk.Widget,
        base_title: str,
        bg_color: str,
        header_color: str,
        **kwargs,
    ) -> None:
        super().__init__(parent, **kwargs)
        self._base_title   = base_title
        self._bg_color     = bg_color
        self._header_color = header_color

        self.articles:        dict[str, list[list[tuple[str, dict]]]] = {}
        self.doc_name:        str = ""
        self._current_key:    str | None = None
        self._has_content:    bool = False
        self._sorted_keys:    list[str] = []
        self._current_index:  int = -1
        self._nav_status_cb:  "callable[[str], None] | None" = None  # ステータスバー更新用CB
        self._file_path:      str | None = None                       # 読み込み済みファイルパス
        self._editing:        bool = False                            # 編集モード中フラグ
        self._editor_bar_visible: bool = False                        # エディタバー表示済みフラグ
        self._edit_state_cb:  "callable[[], None] | None" = None     # 編集状態変更時CB

        self._build()

    # ── UI 構築 ──────────────────────────

    def _build(self) -> None:
        # ヘッダー
        self._header = tk.Label(
            self, text=self._base_title, font=F_BOLD,
            bg=self._header_color, fg="white", padx=12, pady=5, anchor=tk.W,
        )
        self._header.pack(fill=tk.X)

        # ── エディタバー（ファイル読み込み後に header の直後へ pack する）──
        self._editor_frame = tk.Frame(self, bg="#fdf0d0", padx=8, pady=4)
        # ※ _show_editor_bar() が呼ばれるまで pack しない

        self._btn_edit = tk.Button(
            self._editor_frame, text="編集", command=self._enter_edit,
            padx=10, pady=2, bg="#c05010", fg="white", relief=tk.GROOVE,
            font=("MS Gothic", 10),
        )
        self._btn_edit.pack(side=tk.LEFT, padx=2)

        self._btn_save = tk.Button(
            self._editor_frame, text="保存", command=self._save_doc,
            padx=10, pady=2, bg="#2a7a4a", fg="white", relief=tk.GROOVE,
            font=("MS Gothic", 10),
        )
        # 編集中のみ pack する

        self._btn_back = tk.Button(
            self._editor_frame, text="検索に戻る", command=self._back_to_search,
            padx=10, pady=2, relief=tk.GROOVE,
            font=("MS Gothic", 10),
        )
        # 編集中のみ pack する

        # 検索バー
        sf = tk.Frame(self, bg="#ebebeb", padx=8, pady=5)
        sf.pack(fill=tk.X)
        self._search_frame = sf  # 編集モード時の表示/非表示に使用

        tk.Label(sf, text="第", font=("MS Gothic", 11), bg="#ebebeb").pack(side=tk.LEFT)
        self._art_var = tk.StringVar()
        self._art_ent = tk.Entry(sf, textvariable=self._art_var,
                                  font=("MS Gothic", 11), width=6)
        self._art_ent.pack(side=tk.LEFT, padx=(2, 0))

        tk.Label(sf, text="条の", font=("MS Gothic", 11), bg="#ebebeb").pack(side=tk.LEFT)
        self._sub_var = tk.StringVar()
        self._sub_ent = tk.Entry(sf, textvariable=self._sub_var,
                                  font=("MS Gothic", 11), width=5)
        self._sub_ent.pack(side=tk.LEFT, padx=(2, 6))

        tk.Label(sf, text="（枝番なければそのままEnter）",
                 font=("MS Gothic", 8), fg="#888", bg="#ebebeb").pack(side=tk.LEFT)

        tk.Button(
            sf, text="検索", command=self._on_search,
            padx=10, pady=2, bg="#3572b8", fg="white", relief=tk.GROOVE,
        ).pack(side=tk.RIGHT, padx=4)

        # 条文ナビゲーションボタン（▼ 次の条文 / ▲ 前の条文）
        # pack(side=RIGHT) は右から左へ積まれるため、
        # 画面上の並び順は: [key_lbl] [▲] [▼] [検索]
        tk.Button(
            sf, text="▼", command=lambda: self._btn_navigate(+1),
            padx=7, pady=2, relief=tk.GROOVE,
            font=("MS Gothic", 10), bg="#f5f5f5",
        ).pack(side=tk.RIGHT, padx=(1, 0))

        tk.Button(
            sf, text="▲", command=lambda: self._btn_navigate(-1),
            padx=7, pady=2, relief=tk.GROOVE,
            font=("MS Gothic", 10), bg="#f5f5f5",
        ).pack(side=tk.RIGHT, padx=(4, 1))

        self._key_lbl = tk.Label(sf, text="", font=("MS Gothic", 10, "bold"),
                                  bg="#ebebeb", fg="#333")
        self._key_lbl.pack(side=tk.RIGHT, padx=8)

        # Enter キー
        self._art_ent.bind("<Return>", lambda _: self._sub_ent.focus_set())
        self._sub_ent.bind("<Return>", lambda _: self._on_search())

        # テキスト表示エリア
        txt_frame = tk.Frame(self)
        txt_frame.pack(fill=tk.BOTH, expand=True)

        self._txt = tk.Text(
            txt_frame, wrap=tk.WORD, font=F_DEFAULT,
            padx=14, pady=14, bg=self._bg_color,
            state=tk.DISABLED, relief=tk.FLAT, cursor="arrow",
            selectbackground="#b0d0f0",
        )
        sy = tk.Scrollbar(txt_frame, orient=tk.VERTICAL,   command=self._txt.yview)
        sx = tk.Scrollbar(txt_frame, orient=tk.HORIZONTAL, command=self._txt.xview)
        self._txt.config(yscrollcommand=sy.set, xscrollcommand=sx.set)
        sy.pack(side=tk.RIGHT,  fill=tk.Y)
        sx.pack(side=tk.BOTTOM, fill=tk.X)
        self._txt.pack(fill=tk.BOTH, expand=True)

        self._placeholder(f"「{self._base_title}を開く」でファイルを選択してください")

    # ── ファイルセット ────────────────────

    def load_file(self, path: str) -> None:
        """ファイルを解析してパネルをリセット"""
        if self._editing:
            self._exit_edit()
        articles = parse_docx(path)
        self.articles        = articles
        self.doc_name        = os.path.splitext(os.path.basename(path))[0]
        self._file_path      = path
        self._current_key    = None
        self._has_content    = False
        self._sorted_keys    = sorted(articles.keys(), key=_art_sort_key)
        self._current_index  = -1
        self._header.config(
            text=f"{self._base_title}  ―  {os.path.basename(path)}"
        )
        self._key_lbl.config(text="")
        self._placeholder(
            f"読み込み完了 ({len(articles)} 条文)\n\n"
            "条文番号を入力するか ↓↑ キーで条文を切り替えてください"
        )
        self._show_editor_bar()

    # ── 検索 ──────────────────────────────

    def _on_search(self) -> None:
        """検索バーから検索実行"""
        art = self._art_var.get().strip()
        sub = self._sub_var.get().strip()
        if not art:
            messagebox.showwarning("入力エラー", "条文番号を入力してください", parent=self)
            self._art_ent.focus_set()
            return
        key = (
            f"第{normalize_num(art)}条の{normalize_num(sub)}"
            if sub else f"第{normalize_num(art)}条"
        )
        self.show_key(key)

    def show_key(self, key: str) -> None:
        """外部からキーを指定して表示（ESC 一括検索などで使用）"""
        # 検索バーの表示を更新
        m = re.match(r"第(\d+)条(?:の(\d+))?", key)
        if m:
            self._art_var.set(m.group(1))
            self._sub_var.set(m.group(2) or "")
        # _current_index を更新
        if key in self._sorted_keys:
            self._current_index = self._sorted_keys.index(key)
        self._render_key(key)

    def navigate(self, delta: int) -> str | None:
        """delta (+1 or -1) だけ条文を移動して表示。新しいキーを返す。"""
        if not self._sorted_keys:
            return None
        if self._current_index == -1:
            # まだ未選択 → ↓ なら先頭、↑ なら末尾へ
            new_idx = 0 if delta > 0 else len(self._sorted_keys) - 1
        else:
            new_idx = self._current_index + delta
            new_idx = max(0, min(new_idx, len(self._sorted_keys) - 1))
            if new_idx == self._current_index:
                return None  # 端に達しており移動なし
        key = self._sorted_keys[new_idx]
        self.show_key(key)
        return key

    def _btn_navigate(self, delta: int) -> None:
        """▲/▼ ボタン用: このパネルのみ独立してナビゲート"""
        if not self._sorted_keys:
            return
        key = self.navigate(delta)
        if key is not None and self._nav_status_cb is not None:
            direction = "次" if delta > 0 else "前"
            self._nav_status_cb(f"{self._base_title}  {direction}の条文: {key}")

    def _render_key(self, key: str) -> None:
        content = find_article(self.articles, key)
        self._current_key = key
        self._has_content = content is not None
        self._key_lbl.config(text=f"表示: {key}" if content else f"未検出: {key}")
        self._render(content, key)

    # ── レンダリング（リッチテキスト） ──────

    def _render(
        self,
        content: list[list[tuple[str, dict]]] | None,
        key: str,
    ) -> None:
        self._txt.config(state=tk.NORMAL)
        self._txt.delete("1.0", tk.END)

        # 既存の動的タグをすべて削除（"sel" は tkinter 内部タグなので保持）
        for tag in self._txt.tag_names():
            if tag != "sel":
                self._txt.tag_delete(tag)

        # 静的タグ（エラー表示・ヒント・プレースホルダ用）
        self._txt.tag_config("err",  font=F_DEFAULT, foreground="#cc2222")
        self._txt.tag_config("hint", font=F_SMALL,   foreground="#999999")
        self._txt.tag_config("ph",   font=F_DEFAULT, foreground="#aaaaaa")

        if content:
            run_counter = [0]

            def insert_rich_para(rich_para: list[tuple[str, dict]]) -> None:
                """1段落分のリッチテキストを Text ウィジェットに挿入"""
                if not rich_para:
                    self._txt.insert(tk.END, "\n")
                    return
                for text, fmt in rich_para:
                    if not text:
                        continue
                    run_counter[0] += 1
                    tag_name = f"r{run_counter[0]}"

                    # ── フォント ──────────────────────────
                    fn_raw  = fmt.get("font_name") or None
                    fn_win  = (_FONT_MAP.get(fn_raw, fn_raw) if fn_raw
                               else "MS Gothic")
                    fs_raw  = fmt.get("font_size")
                    fs      = max(1, round(fs_raw)) if fs_raw else 11

                    styles: list[str] = []
                    if fmt.get("bold"):
                        styles.append("bold")
                    if fmt.get("italic"):
                        styles.append("italic")
                    font_spec = (
                        (fn_win, fs, " ".join(styles)) if styles
                        else (fn_win, fs)
                    )

                    # ── タグ設定 ──────────────────────────
                    tag_kwargs: dict = {"font": font_spec}
                    clr = fmt.get("color")
                    if clr:
                        tag_kwargs["foreground"] = clr
                    if fmt.get("underline"):
                        tag_kwargs["underline"] = 1
                    if fmt.get("strike"):
                        tag_kwargs["overstrike"] = 1

                    self._txt.tag_config(tag_name, **tag_kwargs)
                    self._txt.insert(tk.END, text, tag_name)

                self._txt.insert(tk.END, "\n")

            for para in content:
                insert_rich_para(para)

        elif self.articles:
            self._txt.insert(tk.END, f"「{key}」は見つかりませんでした。\n\n", "err")
            sample = sorted(self.articles.keys(), key=_art_sort_key)[:15]
            self._txt.insert(tk.END, "【存在する条文（先頭 15 件）】\n", "hint")
            for k in sample:
                self._txt.insert(tk.END, f"    {k}\n", "hint")
        else:
            self._txt.insert(tk.END, "ファイルが読み込まれていません", "ph")

        self._txt.config(state=tk.DISABLED)

    def _placeholder(self, msg: str) -> None:
        self._txt.config(state=tk.NORMAL)
        self._txt.delete("1.0", tk.END)
        for tag in self._txt.tag_names():
            if tag != "sel":
                self._txt.tag_delete(tag)
        self._txt.tag_config("ph", foreground="#aaa", font=F_DEFAULT)
        self._txt.insert(tk.END, f"\n\n  {msg}", "ph")
        self._txt.config(state=tk.DISABLED)

    # ── テキスト出力用 ────────────────────

    def get_export_info(self) -> tuple[str | None, str | None, str | None]:
        """(doc_name, current_key, content_text) を返す。未検索・未検出は None"""
        if not self._has_content or self._current_key is None:
            return None, None, None
        content = find_article(self.articles, self._current_key)
        if content is None:
            return None, None, None
        plain = _article_plain_text(content)
        return self.doc_name, self._current_key, plain

    # ── エディタバー表示 ──────────────────

    def _show_editor_bar(self) -> None:
        """エディタバーを header の直後に一度だけ pack する"""
        if not self._editor_bar_visible:
            self._editor_frame.pack(fill=tk.X, after=self._header)
            self._editor_bar_visible = True

    # ── 編集モード入退 ────────────────────

    @property
    def is_editing(self) -> bool:
        return self._editing

    def _enter_edit(self) -> None:
        """このパネルを編集モードへ切り替える"""
        if not self._file_path:
            messagebox.showwarning("エラー", "先にファイルを読み込んでください", parent=self)
            return
        full_text = self._load_full_text()
        # 検索バーを隠す
        self._search_frame.pack_forget()
        # テキストウィジェットを編集可能にする
        self._txt.config(state=tk.NORMAL, bg="#fffff8", cursor="xterm")
        self._txt.delete("1.0", tk.END)
        for tag in self._txt.tag_names():
            if tag != "sel":
                self._txt.tag_delete(tag)
        self._txt.insert("1.0", full_text)
        # ボタンを切り替える
        self._btn_edit.pack_forget()
        self._btn_save.pack(side=tk.LEFT, padx=2)
        self._btn_back.pack(side=tk.LEFT, padx=(2, 8))
        self._editing = True
        if self._edit_state_cb:
            self._edit_state_cb()
        # 表示していた条文の位置へスクロール
        if self._current_key:
            self._txt.after(0, lambda: self._scroll_to_article(self._current_key))

    def _scroll_to_article(self, key: str) -> None:
        """テキストウィジェット内で key に対応する条文行へスクロールする"""
        m_key = re.match(r"第(\d+)条(?:の(\d+))?", key)
        if not m_key:
            return
        target_art = m_key.group(1)
        target_sub = m_key.group(2)  # 枝番なければ None

        content = self._txt.get("1.0", tk.END)
        for line_no, line in enumerate(content.split("\n"), start=1):
            m = _ART_RE.match(line)
            if m:
                art = normalize_num(m.group("art"))
                sub = normalize_num(m.group("sub")) if m.group("sub") else None
                if art == target_art and sub == target_sub:
                    pos = f"{line_no}.0"
                    self._txt.see(pos)
                    self._txt.mark_set(tk.INSERT, pos)
                    return

    def _exit_edit(self) -> None:
        """編集モードを終了して読み取り専用に戻す（内部用）"""
        self._editing = False
        self._txt.config(state=tk.DISABLED, bg=self._bg_color, cursor="arrow")
        # 検索バーを復元（エディタバーの直後）
        self._search_frame.pack(fill=tk.X, after=self._editor_frame)
        # ボタンを切り替える
        self._btn_save.pack_forget()
        self._btn_back.pack_forget()
        self._btn_edit.pack(side=tk.LEFT, padx=2)
        if self._edit_state_cb:
            self._edit_state_cb()

    def _back_to_search(self) -> None:
        """「検索に戻る」: 保存せずに Viewer モードへ戻る"""
        self._exit_edit()
        # 最後に表示していた条文を再描画
        if self._current_key and self._has_content:
            self._render_key(self._current_key)
        elif self.articles:
            self._placeholder(
                f"読み込み完了 ({len(self.articles)} 条文)\n\n"
                "条文番号を入力するか ↓↑ キーで条文を切り替えてください"
            )

    # ── 保存 ─────────────────────────────

    def _save_doc(self) -> None:
        """「保存」: 上書き保存 or 別名で保存を選択して保存"""
        content = self._txt.get("1.0", tk.END).rstrip("\n")

        answer = messagebox.askyesnocancel(
            "保存方法を選択",
            f"保存方法を選択してください。\n\n"
            f"[はい]      上書き保存  ―  {os.path.basename(self._file_path)}\n"
            f"[いいえ]    別名で保存\n"
            f"[キャンセル] 保存しない",
            parent=self,
        )
        if answer is None:   # キャンセル
            return

        if answer:           # はい → 上書き
            save_path = self._file_path
        else:                # いいえ → 別名
            save_path = filedialog.asksaveasfilename(
                title="別名で保存",
                initialdir=os.path.dirname(self._file_path),
                initialfile=os.path.basename(self._file_path),
                filetypes=[("Word 文書", "*.docx"), ("すべてのファイル", "*.*")],
                defaultextension=".docx",
                parent=self,
            )
            if not save_path:
                return

        try:
            self._write_docx(save_path, content)
        except Exception as exc:
            messagebox.showerror("保存エラー", str(exc), parent=self)
            return

        # 保存成功 → ファイルパスを更新して編集モードを抜ける
        old_key = self._current_key
        self._file_path = save_path
        self._exit_edit()

        # 保存したファイルを再解析
        try:
            articles = parse_docx(save_path)
            self.articles       = articles
            self.doc_name       = os.path.splitext(os.path.basename(save_path))[0]
            self._sorted_keys   = sorted(articles.keys(), key=_art_sort_key)
            self._current_index = -1
            self._header.config(
                text=f"{self._base_title}  ―  {os.path.basename(save_path)}"
            )
        except Exception:
            pass  # 再解析に失敗しても保存自体は完了

        messagebox.showinfo("保存完了", f"保存しました:\n{save_path}", parent=self)

        if old_key:
            self.show_key(old_key)
        else:
            self._placeholder(
                f"保存完了 ({len(self.articles)} 条文)\n\n"
                "条文番号を入力するか ↓↑ キーで条文を切り替えてください"
            )

    # ── ファイル I/O ──────────────────────

    def _load_full_text(self) -> str:
        """docx の全段落をプレーンテキストとして結合して返す"""
        doc = Document(self._file_path)
        return "\n".join(para.text for para in doc.paragraphs)

    def _write_docx(self, path: str, text: str) -> None:
        """テキストを .docx ファイルとして保存（1行 = 1段落）"""
        doc = Document()
        for line in text.split("\n"):
            doc.add_paragraph(line.rstrip("\r"))
        doc.save(path)


# ─────────────────────────────────────────
#  メインウィンドウ
# ─────────────────────────────────────────

class CompareViewer:

    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self.root.title("規約比較ビューワー")
        self.root.geometry("1440x880")
        self._build_ui()
        # 編集状態変更時のコールバックを各パネルに設定
        self._panel1._edit_state_cb = self._on_edit_state_changed
        self._panel2._edit_state_cb = self._on_edit_state_changed
        self.root.bind("<Escape>", lambda _: self._on_escape())
        self.root.bind("<Down>",   lambda _: self._navigate(+1))
        self.root.bind("<Up>",     lambda _: self._navigate(-1))

    def _build_ui(self) -> None:
        # ── ツールバー ──
        bar = tk.Frame(self.root, bg="#e0e0e0", padx=8, pady=6)
        bar.pack(fill=tk.X)

        def _btn(text, cmd, **kw):
            return tk.Button(bar, text=text, command=cmd,
                             padx=9, pady=3, relief=tk.GROOVE, **kw)

        _btn("比較元規約を開く",   self._load1).pack(side=tk.LEFT, padx=3)
        _btn("比較対象規約を開く", self._load2).pack(side=tk.LEFT, padx=3)

        tk.Frame(bar, width=1, bg="#bbb").pack(side=tk.LEFT, fill=tk.Y, padx=10, pady=2)

        self._unified_btn = _btn("両パネル一括検索  [ESC]", self._unified_search,
                                  bg="#3572b8", fg="white")
        self._unified_btn.pack(side=tk.LEFT, padx=3)

        tk.Frame(bar, width=1, bg="#bbb").pack(side=tk.LEFT, fill=tk.Y, padx=10, pady=2)

        _btn("テキスト出力", self._export_text,
             bg="#4a8a30", fg="white").pack(side=tk.LEFT, padx=3)

        self._status = tk.StringVar(value="ファイルを読み込んでください")
        tk.Label(bar, textvariable=self._status, bg="#e0e0e0",
                 fg="#555", font=F_SMALL).pack(side=tk.RIGHT, padx=10)

        # ── 分割ペイン ──
        pw = tk.PanedWindow(self.root, orient=tk.HORIZONTAL,
                             sashwidth=5, sashrelief=tk.FLAT, bg="#bbb")
        pw.pack(fill=tk.BOTH, expand=True, padx=4, pady=(0, 4))

        self._panel1 = DocPanel(pw, "比較元規約",   BG_LEFT,  C_LEFT_H)
        pw.add(self._panel1, minsize=400, stretch="always")

        self._panel2 = DocPanel(pw, "比較対象規約", BG_RIGHT, C_RIGHT_H)
        pw.add(self._panel2, minsize=400, stretch="always")

        # ▲/▼ ボタン押下時にステータスバーを更新するコールバックを設定
        self._panel1._nav_status_cb = self._status.set
        self._panel2._nav_status_cb = self._status.set

        # ── ステータスバー ──
        sb = tk.Frame(self.root, bg="#d0d0d0", height=22)
        sb.pack(fill=tk.X, side=tk.BOTTOM)
        sb.pack_propagate(False)
        tk.Label(sb, textvariable=self._status, bg="#d0d0d0",
                 font=F_SMALL, anchor=tk.W).pack(side=tk.LEFT, padx=8, pady=2)

    # ── 編集状態管理 ──────────────────────

    def _any_editing(self) -> bool:
        """いずれかのパネルが編集中かどうかを返す"""
        return self._panel1.is_editing or self._panel2.is_editing

    def _on_edit_state_changed(self) -> None:
        """パネルの編集状態が変わったときに呼ばれる"""
        editing = self._any_editing()
        state = tk.DISABLED if editing else tk.NORMAL
        self._unified_btn.config(state=state)
        if editing:
            self._status.set("編集モード中 — 検索・ナビゲーション機能は抑止されています")

    def _on_escape(self) -> None:
        """ESC キー: 編集中は無視し、それ以外は一括検索を実行"""
        if self._any_editing():
            return
        self._unified_search()

    # ── ファイル読み込み ────────────────

    def _load1(self) -> None:
        path = filedialog.askopenfilename(
            title="比較元規約ファイルを選択",
            filetypes=[("Word 文書", "*.docx"), ("すべてのファイル", "*.*")],
        )
        if not path:
            return
        try:
            self._panel1.load_file(path)
            self._status.set(
                f"比較元規約 読み込み完了  ({len(self._panel1.articles)} 条文検出)"
            )
        except Exception as exc:
            messagebox.showerror("読み込みエラー", str(exc))

    def _load2(self) -> None:
        path = filedialog.askopenfilename(
            title="比較対象規約ファイルを選択",
            filetypes=[("Word 文書", "*.docx"), ("すべてのファイル", "*.*")],
        )
        if not path:
            return
        try:
            self._panel2.load_file(path)
            self._status.set(
                f"比較対象規約 読み込み完了  ({len(self._panel2.articles)} 条文検出)"
            )
        except Exception as exc:
            messagebox.showerror("読み込みエラー", str(exc))

    # ── カーソルキーナビゲーション ─────────

    def _navigate(self, delta: int) -> None:
        """↓/↑ キーで両パネルの条文を移動。Entry にフォーカスがある場合はスキップ。"""
        if self._any_editing():
            return  # 編集モード中はナビゲーションを抑止
        focused = self.root.focus_get()
        if isinstance(focused, tk.Entry):
            return  # 入力フィールド操作中は矢印キーを横取りしない

        key1 = self._panel1.navigate(delta)
        key2 = self._panel2.navigate(delta)

        displayed = key1 or key2
        if displayed:
            direction = "次" if delta > 0 else "前"
            self._status.set(f"{direction}の条文: {displayed}")

    # ── 一括検索（ESC） ─────────────────

    def _unified_search(self) -> None:
        """ESC キー: 両パネルに同じ条文番号で検索"""
        if self._any_editing():
            return  # 編集モード中は検索を抑止
        if not self._panel1.articles and not self._panel2.articles:
            messagebox.showinfo("お知らせ", "先にファイルを読み込んでください")
            return
        dlg = SearchDialog(self.root)
        self.root.wait_window(dlg)
        if dlg.result:
            art, sub = dlg.result
            key = (
                f"第{normalize_num(art)}条の{normalize_num(sub)}"
                if sub else f"第{normalize_num(art)}条"
            )
            self._panel1.show_key(key)
            self._panel2.show_key(key)
            self._status.set(f"両パネルで {key} を検索しました")

    # ── テキスト出力 ───────────────────

    def _export_text(self) -> None:
        """表示中の条文をテキストファイルに出力"""
        name1, key1, text1 = self._panel1.get_export_info()
        name2, key2, text2 = self._panel2.get_export_info()

        if not text1 and not text2:
            messagebox.showinfo(
                "出力できません",
                "出力する条文がありません。\n先に条文を検索してください。",
            )
            return

        save_dir = filedialog.askdirectory(title="保存先フォルダを選択")
        if not save_dir:
            return

        folder_key = key1 or key2
        folder_key_safe = re.sub(r'[\\/:*?"<>|]', "_", folder_key)
        folder_path = os.path.join(save_dir, folder_key_safe)

        try:
            os.makedirs(folder_path, exist_ok=True)
        except OSError as e:
            messagebox.showerror("フォルダ作成エラー", f"フォルダを作成できませんでした:\n{e}")
            return

        saved: list[str] = []
        errors: list[str] = []

        for doc_name, key, text in [
            (name1, key1, text1),
            (name2, key2, text2),
        ]:
            if not text or not doc_name or not key:
                continue
            key_safe = re.sub(r'[\\/:*?"<>|]', "_", key)
            fname = f"{doc_name}_{key_safe}.txt"
            fpath = os.path.join(folder_path, fname)
            try:
                with open(fpath, "w", encoding="utf-8") as f:
                    f.write(text)
                saved.append(fname)
            except OSError as e:
                errors.append(f"{fname}: {e}")

        if saved:
            msg = (
                f"フォルダ「{folder_key_safe}」に以下のファイルを保存しました:\n\n"
                + "\n".join(saved)
            )
            if errors:
                msg += "\n\nエラー:\n" + "\n".join(errors)
            messagebox.showinfo("出力完了", msg)
            self._status.set(
                f"テキスト出力完了: {folder_key_safe}/ → {', '.join(saved)}"
            )
        else:
            messagebox.showerror("出力失敗", "\n".join(errors))


# ─────────────────────────────────────────
#  一括検索ダイアログ（ESC キー用）
# ─────────────────────────────────────────

class SearchDialog(tk.Toplevel):

    def __init__(self, parent: tk.Tk) -> None:
        super().__init__(parent)
        self.result: tuple[str, str] | None = None
        self.title("両パネル一括検索")
        self.geometry("420x260")
        self.resizable(False, False)
        self.transient(parent)
        self.grab_set()

        self.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width()  - 420) // 2
        y = parent.winfo_y() + (parent.winfo_height() - 260) // 2
        self.geometry(f"+{x}+{y}")

        self._build()
        self.bind("<Escape>", self._cancel)
        self._art_ent.focus_set()

    def _build(self) -> None:
        outer = tk.Frame(self, padx=30, pady=24)
        outer.pack(fill=tk.BOTH, expand=True)

        tk.Label(outer, text="表示する条文番号を入力してください。",
                 font=("MS Gothic", 11)).grid(
            row=0, column=0, columnspan=4, sticky=tk.W, pady=(0, 6))

        tk.Label(outer, text="第", font=("MS Gothic", 13)).grid(
            row=1, column=0, sticky=tk.W)
        self._art_var = tk.StringVar()
        self._art_ent = tk.Entry(outer, textvariable=self._art_var,
                                  font=("MS Gothic", 13), width=8)
        self._art_ent.grid(row=1, column=1, sticky=tk.W, padx=4)
        tk.Label(outer, text="条", font=("MS Gothic", 13)).grid(
            row=1, column=2, sticky=tk.W)

        tk.Label(outer, text="枝番号を入力してください。（なければそのままEnterを押す）",
                 font=("MS Gothic", 11)).grid(
            row=2, column=0, columnspan=4, sticky=tk.W, pady=(20, 6))

        tk.Label(outer, text="の", font=("MS Gothic", 13)).grid(
            row=3, column=0, sticky=tk.W)
        self._sub_var = tk.StringVar()
        self._sub_ent = tk.Entry(outer, textvariable=self._sub_var,
                                  font=("MS Gothic", 13), width=8)
        self._sub_ent.grid(row=3, column=1, sticky=tk.W, padx=4)
        tk.Label(outer, text="（省略可）", font=("MS Gothic", 9), fg="#999").grid(
            row=3, column=2, columnspan=2, sticky=tk.W)

        bf = tk.Frame(outer)
        bf.grid(row=4, column=0, columnspan=4, sticky=tk.E, pady=(24, 0))
        tk.Button(bf, text="キャンセル", command=self._cancel,
                  padx=10, pady=4).pack(side=tk.RIGHT, padx=(8, 0))
        tk.Button(bf, text="検索", command=self._ok,
                  padx=16, pady=4, bg="#3572b8", fg="white",
                  relief=tk.GROOVE).pack(side=tk.RIGHT)

        self._art_ent.bind("<Return>", lambda _: self._sub_ent.focus_set())
        self._sub_ent.bind("<Return>", lambda _: self._ok())

    def _ok(self) -> None:
        art = self._art_var.get().strip()
        sub = self._sub_var.get().strip()
        if not art:
            messagebox.showwarning("入力エラー", "条文番号を入力してください", parent=self)
            self._art_ent.focus_set()
            return
        self.result = (art, sub)
        self.destroy()

    def _cancel(self, _event=None) -> None:
        self.destroy()


# ─────────────────────────────────────────
#  エントリーポイント
# ─────────────────────────────────────────

if __name__ == "__main__":
    root = tk.Tk()
    CompareViewer(root)
    root.mainloop()
